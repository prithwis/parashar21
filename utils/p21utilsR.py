#Copyright (c) 2022, Prithwis Mukerjee
#All rights reserved.
#
#This source code is licensed under the GNU GPL v3.0 -style license found in the
#LICENSE file in the root directory of this source tree. 
# --------------------------------------------------
# Global Variables
import p21
#import p21swe
import p21utils
import p21YogInfo
# --------------------------------------------------
#Utility functions 

import pandas as pd
import dateutil
import matplotlib.pyplot as plt
import math
import math
import numbers
#import collections
#import string_utils

import json

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime
from datetime import timedelta

#import swisseph as swe
import pytz



#
# --------------------------------------------------
#
# Generates the text that appears in a chart image
#
def generateChartTxt():
    global txt
    # .........  using this function for some additional, common tasks
    #
    if p21.ChartType == 'Rashi':
        p21.ChartColour = 'white'      # instead of orange
        p21.ChartFile = 'RashiChart.png'
    else:
        p21.ChartColour = 'white'      # instead of olive
        p21.ChartFile = 'NavamsaChart.png'
    # ......    
    txt = ['']*13
    for g,r in p21.GRashiN.items():
        if r == 1:
            txt[1]= txt[1]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 2:
            txt[2]= txt[2]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 3:
            txt[3]= txt[3]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 4:
            txt[4]= txt[4]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 5:
            txt[5]= txt[5]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 6:
            txt[6]= txt[6]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 7:
            txt[7]= txt[7]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 8:
            txt[8]= txt[8]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 9:
            txt[9]= txt[9]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 10:
            txt[10]= txt[10]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 11:
            txt[11]= txt[11]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 12:
            txt[12]= txt[12]+' '+g+('/R' if p21.GRet[g] else '')
            
    for i in range(1,13):
        if len(txt[i]) == 0:
            txt[i] = '*'
        if p21.pName == p21.gName:
            if p21.ChartType == 'Rashi':
                txt[i] = txt[i].replace('La','RS')
            else:
                txt[i] = 'NA'
#print(txt[1],txt[2],txt[3],txt[4],txt[5],txt[6],txt[7],txt[8],txt[9],txt[10],txt[11],txt[12],)

# --------------------------------------------------
# Plots the chart in the Bengal Format
#
def R12B_drawChart_Bengal():

    #ChartFile = p21.pName+p21.ChartType+'.png'
    #ChartFile = 'CurrentChart.png'
    generateChartTxt()

    #id = ChartType+'\n'+pName+'\n'+pDate+'\n'+pTime+'\n'+pPlace
    id = p21.ChartType+'\n'+p21.pName
    '''
    if p21.ChartType == 'Rashi':
        ChartColour = 'orange'
        ChartFile = 'RashiChart.png'
    else:
        ChartColour = 'olive'
        ChartFile = 'NavamsaChart.png'
    '''
        
    #plt.figure(figsize=(7,7))
    #plt.figure(figsize=(7,7),facecolor=ChartColour)
    plt.figure(figsize=(7,7),facecolor=p21.ChartColour)
    
    plt.axis('off')

    # draw vertical line 
    plt.plot([30, 30], [0, 90], 'k-', lw=2)
    plt.plot([60, 60], [0, 90], 'k-', lw=2)

    # draw horizontal line 
    plt.plot([0, 90], [30, 30], 'k-', lw=2)
    plt.plot([0, 90], [60, 60], 'k-', lw=2)

    #draw diagonal lines
    plt.plot([60,90],[60,90], 'k-', lw=2)
    plt.plot([0,30],[90,60], 'k-', lw=2)
    plt.plot([0,30],[90,60], 'k-', lw=2)
    plt.plot([0,30],[0,30], 'k-', lw=2)
    plt.plot([60,90],[30,0], 'k-', lw=2)

    plt.text(32, 38, id, fontsize=12)
    plt.text(0, 95, p21.cID, fontsize=12)       # unique chartID

    plt.text(32, 82, txt[1], fontsize=12)
    plt.text(8, 82, txt[2], fontsize=12)
    plt.text(2, 62, txt[3], fontsize=12)
    plt.text(2, 45, txt[4], fontsize=12)
    plt.text(2, 25, txt[5], fontsize=12)
    plt.text(8, 5, txt[6], fontsize=12)
    plt.text(32, 5, txt[7], fontsize=12)
    plt.text(62, 5, txt[8], fontsize=12)
    plt.text(68, 25, txt[9], fontsize=12)
    plt.text(68, 45, txt[10], fontsize=12)
    plt.text(68, 62, txt[11], fontsize=12)
    plt.text(62, 82, txt[12], fontsize=12)

    
    #plt.savefig("CurrentChart.png", bbox_inches='tight')
    #plt.savefig(ChartFile, bbox_inches='tight')
    plt.savefig(p21.ChartFile, bbox_inches='tight')
    #plt.show()

#Draw Chart as Per South India Style
def R12B_drawChart_South():

    #ChartFile = p21.pName+p21.ChartType+'.png'
    #ChartFile = 'CurrentChart.png'
    generateChartTxt()

    #id = ChartType+'\n'+pName+'\n'+pDate+'\n'+pTime+'\n'+pPlace
    id = p21.ChartType+'\n'+p21.pName
    '''
    if p21.ChartType == 'Rashi':
        ChartColour = 'orange'
        ChartFile = 'RashiChart.png'
    else:
        ChartColour = 'olive'
        ChartFile = 'NavamsaChart.png'
    '''
        
    #plt.figure(figsize=(7,7))
    plt.figure(figsize=(7,7),facecolor=p21.ChartColour)
    
    plt.axis('off')

    plt.text(25, 35, id, fontsize=12)
    plt.text(0, 95, p21.cID, fontsize=12)       # unique chartID

    # draw vertical line 
    plt.plot([0, 0], [0, 90], 'k-', lw=2)
    plt.plot([22.5, 22.5], [0, 90], 'k-', lw=2)

    plt.plot([45, 45], [0, 22.5], 'k-', lw=2)
    plt.plot([45, 45], [67.5, 90], 'k-', lw=2)


    plt.plot([67.5, 67.5], [0, 90], 'k-', lw=2)
    plt.plot([90, 90], [0, 90], 'k-', lw=2)

    # draw horizontal line

    plt.plot([0, 90], [0, 0], 'k-', lw=2)
    plt.plot([0, 90], [22.5, 22.5], 'k-', lw=2)

    plt.plot([0, 22.5], [45, 45], 'k-', lw=2)
    plt.plot([67.5, 90], [45, 45], 'k-', lw=2)
    
    plt.plot([0, 90], [67.5, 67.5], 'k-', lw=2)
    plt.plot([0, 90], [90, 90], 'k-', lw=2)
    
    #draw diagonal lines
    plt.plot([22.5, 27.5], [85, 90], 'k-', lw=2)

    # graha names

    plt.text(0.5, 80, txt[12], fontsize=15)
    plt.text(23, 73, txt[1], fontsize=15)
    plt.text(45.5, 80, txt[2], fontsize=15)
    plt.text(68, 73, txt[3], fontsize=15)

    plt.text(68, 55, txt[4], fontsize=15)
    plt.text(68, 35, txt[5], fontsize=15)

    plt.text(68, 13, txt[6], fontsize=15)
    plt.text(45.5, 6, txt[7], fontsize=15)
    plt.text(23, 13, txt[8], fontsize=15)
    plt.text(0.5, 6, txt[9], fontsize=15)
   
    plt.text(0.5, 55, txt[11], fontsize=15)
    plt.text(0.5, 35, txt[10], fontsize=15)

    
    #plt.savefig("CurrentChart.png", bbox_inches='tight')
    plt.savefig(p21.ChartFile, bbox_inches='tight')
    #plt.show()


#Draw Chart as Per North India Style
def R12B_drawChart_North():

    #ChartFile = p21.pName+p21.ChartType+'.png'
    #ChartFile = 'CurrentChart.png'
    generateChartTxt()
    
    # for North India, we need to rotate the texts
    LagnaRashiNumber = p21.GRashiN['La']
    ntxt = ['']*13
    j = LagnaRashiNumber
    for i in range(1,13):
        ntxt[i] = str(j)+"|"+txt[j]
        j = j+1
        if j > 12:
            j = 1
    
    #text has to be made vertical for these bhavs!     
    for i in [3,5,9,11]:
        ntxt[i] = ntxt[i].replace(' ','\n')

    #id = ChartType+'\n'+pName+'\n'+pDate+'\n'+pTime+'\n'+pPlace
    id = p21.ChartType+' | '+p21.pName
    #id = '25-35'
    '''
    if p21.ChartType == 'Rashi':
        ChartColour = 'orange'
        ChartFile = 'RashiChart.png'
    else:
        ChartColour = 'olive'
        ChartFile = 'NavamsaChart.png'
    '''
    #print(p21.GRashiN)    
    #plt.figure(figsize=(7,7))
    plt.figure(figsize=(7,7),facecolor=p21.ChartColour)
    
    plt.axis('off')

    plt.text(0, -5, id, fontsize=12)
    plt.text(0, 95, p21.cID, fontsize=12)       # unique chartID
    
    
    
    # draw vertical line 
    plt.plot([0, 0], [0, 90], 'k-', lw=2)  # Left Vertical   [x1,x2], [y1,y2]
    plt.plot([90, 90], [0, 90], 'k-', lw=2) # Right Vertical
    
    # draw horizontal lines
    plt.plot([0, 90], [90, 90], 'k-', lw=2)  # Top Horizontal   [x1,x2], [y1,y2]
    plt.plot([0, 90], [0, 0], 'k-', lw=2)  # Bottom Horizontal   [x1,x2], [y1,y2]
    plt.plot([42, 48], [86, 86], 'k-', lw=2)  # Lagna Horizontal   [x1,x2], [y1,y2]
    
    
    # Diagonal
    plt.plot([0, 90], [0, 90], 'k-', lw=2)  # Bottom Left Top Right   [x1,x2], [y1,y2]
    plt.plot([0, 90], [90, 0], 'k-', lw=2)  # Top Left Bottom Right   [x1,x2], [y1,y2]
    plt.plot([0, 45], [45, 90], 'k-', lw=2) # upward
    plt.plot([45, 90], [0, 45], 'k-', lw=2)  # upward
    plt.plot([0, 45], [45, 0], 'k-', lw=2)
    plt.plot([45, 90], [90, 45], 'k-', lw=2)
    
    
    plt.text(40, 55, p21.RashiN2A(LagnaRashiNumber), fontsize=14)
    
    plt.text(27.5, 67.5, ntxt[1], fontsize=12) # 1
    plt.text(27.5, 22.5, ntxt[7], fontsize=12)  # 7
    
    plt.text(5, 45, ntxt[4], fontsize=12) # 4
    plt.text(50, 45, ntxt[10], fontsize=12)  # 10
    
    plt.text(7, 85, ntxt[2], fontsize=12) # 2
    plt.text(53, 85, ntxt[12], fontsize=12)  # 12
    
    plt.text(7, 3, ntxt[6], fontsize=12) # 6
    plt.text(53, 3, ntxt[8], fontsize=12)  # 8
    
    plt.text(1, 55, ntxt[3], fontsize=12) # 3
    plt.text(1, 10, ntxt[5], fontsize=12)  # 5
    
    plt.text(82, 10, ntxt[9], fontsize=12) # 9
    plt.text(82, 55, ntxt[11], fontsize=12)  # 11
    
    
    plt.savefig(p21.ChartFile, bbox_inches='tight')
    #plt.show()
    #print('north')
#
# Functions to convert lists to strings for transfering to MS Word docs
#

def R13C_ListPositions(text,L):
    # 
    L1 = L[1:]
    D = ''
    for c,e in enumerate(L1,1):
        if e :
            D = D+' '+str(c)+' ' 
    if len(D)>0:
        return(text+D)
    else:
        return(text+'<none>')    
    #return(D)

def R13B_ListPositions(text,L):
    # 
    L1 = L[1:]
    D = text+' : '
    for c,e in enumerate(L1,1):
        D = D+' '+str(c)+':'+str(e)+' '
    return(D)
        
def R13D_ListContents(text,L):
    # 
    #L1 = L[1:]
    D = text+'\t'
    for c,e in enumerate(L,1):
        D = D+str(e)+'\t'
    return(D)

def R13A_ShowTrueDict(desc,tDict):
    # creates a string with Keys that have Value True
    
    D = ''
    for (k,v) in tDict.items():
        if v :
            D = D+' '+k
    if len(D)>0:
        return(desc+D)
    else:
        return(desc+'<none>')
    #return(D)


def R30_LocateYogs():
    p21.yogsFound = []
    for yogName in list(p21YogInfo.yogText.keys()):                                             # Yog conditions picked up from p21YogInfo.py
        objList = list(p21.kollection.find(p21YogInfo.yogCond[yogName],{"pid.name" : 1,"_id":0}))   # Get names of people whose charts have these Yogs
        personList = list(map(lambda e:e['pid']['name'],objList))                               # Convert list of json object to list of string
        if p21.pName in personList:
            p21.yogsFound.append(yogName)

def R511_parseChartData(c):

    #print(c)
    #print('parse',c['maleficG'])
    p21.GLon = c['GLon']
    p21.GRet = c['GRet']
    pid = c['pid']
    p21.pName = pid['name']
    p21.cID = pid['cID']
    p21.pTags = pid['tags']
    p21.GrahaLordBhav = c['GrahaLordBhav']
    p21.Lord = c['Lord']
    p21.GrahaBhava = c['GrahaBhava']
    p21.LordBhav = c['LordBhav']
    p21.exaltG = c['exaltG']
    p21.exaltL = c['exaltL']
    p21.debilG = c['debilG']
    p21.debilL = c['debilL']
    p21.inFriendG = c['inFriendG']
    p21.inFriendL = c['inFriendL']
    p21.inEnemyG = c['inEnemyG']
    p21.inEnemyL = c['inEnemyL']
    
    p21.GAspects2 = c['GAspects2']
    p21.GAspectedBy2 = c['GAspectedBy2']
    p21.BAspectedBy2 = c['BAspectedBy2']
    p21.BAspectedByBL2 = c['BAspectedByBL2']
    
    p21.GConjunctsG2 = c['GConjunctsG2']
    p21.BLConjunctsG2 = c['BLConjunctsG2']
    p21.BLConjunctsBL2 = c['BLConjunctsBL2']
    p21.beneficG = c['beneficG']
    p21.maleficG = c['maleficG']
    

def R01_CreateReportDoc(cqs,pS,repStyle = 'MultiChart'):
      
    now = datetime.now(pytz.timezone('Asia/Kolkata'))
    p21.document = Document()
    
    style = p21.document.styles['Normal']           # Changes the Normal style throuhout the document
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10)                              # Increasing the font size will cause reports to overflow to next page
    
    
    section = p21.document.sections[0]
    header = section.header
    footer = section.footer
    
    header01 = header.paragraphs[0]
    header01.text = "卐 Parashar21 | "+now.strftime("%d %b %Y")
    #header01.text = "卐 Parashar21  "

    
    footer01 = footer.paragraphs[0]
    #footer01.text = "Printed on : "
    footer01.text = "http://parashar21.blogspot.com | https://github.com/prithwis/parashar21"

    if repStyle != 'MultiChart':
        return
    
    #heading_1 = p21.pName +" >>> "+p21.ChartType
    #heading_1 = ChartType+" Chart of "+pName
    
    heading_1 = "Charts of Interest"  #p21.ReportFile
    p21.document.add_heading(heading_1, 0)
    heading_2 = cqs                         # current query string
    p21.document.add_heading(heading_2, level=3)
    p21.document.add_paragraph(pS)          # print count status
    
    p21.document.add_picture('./Saraswati.png', width=Inches(4.25))
    last_paragraph = p21.document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    threeLineGap = '\n\n\n'
    #KNRText = 'based on the book\nLearn Hindu Astrology Easily\nK N Rao & K Ashu Rao\nThird Edition\nVani Publications\nDelhi'
    #para = p21.document.add_paragraph(KNRText)
    #para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p21.document.add_paragraph(threeLineGap)
    #p21.document.add_paragraph(threeLineGap)
    
    gyan = "Astrology is the science of CORRELATIONS\nAstrology - An Application of Data Science \nhttps://bit.ly/pmastro"
    para = p21.document.add_paragraph(gyan)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p21.document.add_paragraph('______________________________________________________________________________')
    
    p21.document.add_picture('./p21logo.png', width=Inches(1.0))
    last_paragraph = p21.document.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    
    
    
    p21.document.add_page_break()
    #return document
    
# --------------------------------------------------



def R512_FormatPage(repStyle = 'MultiChart'):

    
       
    #p21.document.add_page_break()
    p2 = p21.document.add_paragraph()
    run_1 = p2.add_run()
    run_1.add_picture('./RashiChart.png', width=Inches(3.0))
    if p21.pName != p21.gName: 
        run_2 = p2.add_run()
        run_2.add_picture('./NavamsaChart.png', width=Inches(3.0))
    # ----------------------------------------------------------------------------------------------------End of Paragraph    
    #p1 = p21.document.add_paragraph(p21.pName)                                             # skip printing name
    
    
    #cPara = R13A_ShowTrueDict('Retrograde Graha',p21.GRet)+'\n'                          # Show Grahas that are retrograde
    #cPara = ' '.join(p21.pTags)+'\n'
    cPara = 'Details of '+p21.AnalysisType.upper()+' chart\n'
    cPara = cPara+'Tags : '
    cPara = cPara+json.dumps(p21.pTags)
    #cPara = cPara+json.dumps(p21.pTags)+'\n'
    cPara = cPara+'\n.....................................'+'\n'
    cPara = cPara+R13B_ListPositions('Lord of ',p21.Lord)                                # Show Lords
    cPara = cPara+"\n"
    cPara = cPara+'Graha Lord of \n'+json.dumps(p21.GrahaLordBhav).replace('"','')                         # Show the Bhavs of whicha Graha is Lord
    
    cPara = cPara+R13B_ListPositions('\nLord in Bhava',p21.LordBhav)
    p21.document.add_paragraph(cPara)
    
    table = p21.document.add_table(rows=1, cols=2)
    cell1 = table.cell(0, 0)
    cell2 = table.cell(0, 1)
    
    cPara1 = 'Graha Status\n'
    cPara1 = cPara1+R13A_ShowTrueDict('Exalted     : ',p21.exaltG)+'\n'
    cPara1 = cPara1+R13A_ShowTrueDict('Debilitated : ',p21.debilG)+'\n'
    cPara1 = cPara1+R13A_ShowTrueDict('in Friend   : ',p21.inFriendG)+'\n'
    cPara1 = cPara1+R13A_ShowTrueDict('in Enemy    : ',p21.inEnemyG)+'\n'
    cPara1 = cPara1+R13A_ShowTrueDict('Benefics    : ',p21.beneficG)+'\n'
    cPara1 = cPara1+R13A_ShowTrueDict('Malefics    : ',p21.maleficG)
    cell1.add_paragraph(cPara1)
    # ----------------------------------------------------------------------------------------------------End of Paragraph        
    
    
    cPara2 = 'Lord Status\n'
    cPara2 = cPara2+R13C_ListPositions('Exalted     : ',p21.exaltL)+'\n'
    cPara2 = cPara2+R13C_ListPositions('Debilitated : ',p21.debilL)+'\n'
    cPara2 = cPara2+R13C_ListPositions('In Friend   : ',p21.inFriendL)+'\n'
    cPara2 = cPara2+R13C_ListPositions('In Enemy    : ',p21.inEnemyL)
    cell2.add_paragraph(cPara2)
    # ----------------------------------------------------------------------------------------------------End of Paragraph        
    
    cPara = 'Graha Aspects ........................................................'
    txt[0] = json.dumps(p21.GAspects2).replace('"','')
    #cPara = cPara+'\nAspects🤓 '+json.dumps(p21.GAspects2).replace('"','')   
    txt[1] = json.dumps(p21.GAspectedBy2).replace('"','')
    #cPara = cPara+'\nAspected BY 👀'+json.dumps(p21.GAspectedBy2).replace('"','') 
    for i in [0,1]:
        if p21.pName == p21.gName:
            txt[i] = txt[i].replace('La','RS')
    cPara = cPara+'\nAspects🤓 '+txt[0]
    cPara = cPara+'\nAspected BY 👀'+txt[1] 
    #p21.document.add_paragraph(cPara)
    
    #print(p21.GAspects2)
    #print(p21.GAspectedBy2)
    cPara = cPara+'\nBhav Aspect ..................................................'
    #sorted_BAspectedBy2={}
    #print(p21.BAspectedBy2)
    T1 = {}
    for k,v in p21.BAspectedBy2.items():
        T1[int(k)] =v
    T2 = dict(sorted(T1.items()))
    #print(T2)
    txt[0] = json.dumps(T2).replace('"','')
    if p21.pName == p21.gName:
        txt[0] = txt[0].replace('La','RS')
    
    
    #cPara = cPara+'\nAspected BY Graha 👀'+json.dumps(T2).replace('"','')  
    cPara = cPara+'\nAspected BY Graha 👀'+txt[0]  
    
    T1 = {}
    for k,v in p21.BAspectedByBL2.items():
        T1[int(k)] =v
    T2 = dict(sorted(T1.items()))
    #print(T2)
    
    
    txt[0] = json.dumps(T2).replace('"','')
    #print('QQ',txt[0])
    if p21.pName == p21.gName:
        txt[0] = txt[0].replace('La','RS')
    #print('QQ',txt[0])
    cPara = cPara+'\nAspected BY Lords 👀'+txt[0]  
    #cPara = cPara+'\nAspected BY Lords 👀'+json.dumps(T2).replace('"','')  
    #p21.document.add_paragraph(cPara)

    cPara = cPara+'\nConjuncts  ...................................................'
    txt[0] = json.dumps(p21.GConjunctsG2).replace('"','')
    txt[1] = json.dumps(p21.BLConjunctsG2).replace('"','')
    txt[2] = json.dumps(p21.BLConjunctsBL2).replace('"','')
    for i in [0,1,2]:
        if p21.pName == p21.gName:
            txt[i] = txt[i].replace('La','RS')
    
    cPara = cPara+'\nGraha Graha '+txt[0]  
    cPara = cPara+'\nLord Graha '+txt[1] 
    cPara = cPara+'\nLord Lord '+txt[2]
    #cPara = cPara+'\nGraha Graha '+json.dumps(p21.GConjunctsG2).replace('"','')  
    #cPara = cPara+'\nLord Graha '+json.dumps(p21.BLConjunctsG2).replace('"','')  
    #cPara = cPara+'\nLord Lord '+json.dumps(p21.BLConjunctsBL2).replace('"','')
    #p21.document.add_paragraph(cPara)
    
    #cPara = cPara+'\nVocation Tags : '
    #cPara = cPara+json.dumps(p21.pTags)+'\n'
    cPara = cPara+'\nYogs Found : '
    if len(p21.yogsFound) > 0:                      # Yogs Found 
        for y in p21.yogsFound:
            cPara = cPara+y+' '
    p21.document.add_paragraph(cPara)
    # ----------------------------------------------------------------------------------------------------End of Paragraph        
    # Dasha Printing
    #
    #if p21.printDasha:

    # ----------------------------------------------------------------------------------------------------End of Paragraph        
    if (repStyle != 'SingleChart'):
        p21.document.add_page_break()
        

def R512_FormatPage2():
#
    # Ashtakvarga Printing -------------------- ONLY for single chart scenario
    #
    AshtakVarga = p21utils.SarvaAshtakVarga()
    s8v = AshtakVarga[0]
    
    print("p21.GRashiN['La']",p21.GRashiN['La'],'1st Bhav', p21.RashiN2A(p21.GRashiN['La']))
    #cPara = cPara+'\n'+'Ashtakvarga Points for Bhavs starting from '+p21.RashiN2A(p21.GRashiN['La'])
    cPara = '\n'+'Ashtakvarga Points for Bhavs starting from '+p21.RashiN2A(p21.GRashiN['La'])
    s8vB = AshtakVarga[2]
    for i in range(7):
        #print(p21.Graha[i],s8vB[i])
        cPara = cPara + R13D_ListContents('\n'+p21.Graha[i],s8vB[i])
    
    bhavSum = AshtakVarga[3]
    
    
    cPara = cPara + R13D_ListContents('\n +',bhavSum)
    bhavSplit = AshtakVarga[4]
    cPara = cPara + R13D_ListContents('\n Ban 1+5+9:Sev 2+6+10:Pos 3+7+11:Gha 4+8+12:CHK',bhavSplit)
    
    
    para009 = p21.document.add_paragraph(cPara)
    #para001.style = p21.document.styles['Normal']
    para009_format = para009.paragraph_format
    para009_tab_stops = para009_format.tab_stops
    para_tab_stop = para009_tab_stops.add_tab_stop(Inches(0.25))
    # ----------------------------------------------------------------------------------------------------End of Paragraph
    #if repStyle == 'SingleChart':
    if p21.pName == p21.gName:
        cPara = 'Antardasha -----------------\n'
        today = datetime.now()
        #print(today)
        count = 0
        for k1, v1 in p21.VimDasha.items():
            if (today < datetime.strptime(v1['End'],"%d %b %Y")) and (count < 2):
                if (count == 0):
                    #print('Current Dasha', count, k1)
                    cPara = cPara+'Current Dasha '+k1+'\n'
                else:
                    #print('Next Dasha',count
                    cPara = cPara+'Next Dasha \n'
                #print('count', count)
                for k2, v2 in v1.items():
                    if (type(v2) is dict):
                        print(k1,' Dasha ends ',v1['End'],' | ',k2,' Antardasha ends ',v2['End'])
                        cPara = cPara+k1+'\tDasha end\t'+v1['End']+'\t| '+k2+'\tAntardasha ends\t'+v2['End']+'\n'
                    #Error = datetime.strptime(v1['End'],"%d %b %Y") - datetime.strptime(v2['End'],"%d %b %Y")
                    #print('----------- Error of ', Error.days,' days in ',v1['Duration'],' day Dasha')
                count = count+1
        cPara = cPara+'---------------------\n'
    else:
        cPara = 'Mahadasha -----------------\n'
        for k1, v1 in p21.VimDasha.items():
            #print(k1,' Dasha ends on ',v1['End'])
            cPara = cPara+k1+'\tDasha ends on\t'+v1['End']+'\n'
        cPara = cPara+'---------------------\n'
    
    p21.document.add_paragraph(cPara)

#---------------------------------------------------------------------------------------------------
# Dasha Printing Functions

def printVimSottari():
  for k1, v1 in p21.VimDasha.items():
    #print('outer')
    for k2, v2 in v1.items():
      if (type(v2) is dict):
        print(k1,' Dasha ends on ',v1['End'],' | ',k2,' Antardasha ends on ',v2['End'])
    Error = datetime.strptime(v1['End'],"%d %b %Y") - datetime.strptime(v2['End'],"%d %b %Y")
    print('----------- Error of ', Error.days,' days in ',v1['Duration'],' day Dasha')
#
#-------------
#  
def printMahaDasha():
  for k1, v1 in p21.VimDasha.items():
    print(k1,' Dasha ends on ',v1['End'])
#
#-------------
#
def printAntarDasha():
    #today = datetime.now(pytz.timezone('Asia/Calcutta'))
    today = datetime.now()
    print(today)
    count = 0
    for k1, v1 in p21.VimDasha.items():
        if (today < datetime.strptime(v1['End'],"%d %b %Y")) and (count < 2):
            if (count == 0):
                print('Current Dasha', count, k1)
            else:
                print('Next Dasha',count)
            #print('count', count)
            for k2, v2 in v1.items():
                if (type(v2) is dict):
                    print(k1,' Dasha ends ',v1['End'],' | ',k2,' Antardasha ends ',v2['End'])
                #Error = datetime.strptime(v1['End'],"%d %b %Y") - datetime.strptime(v2['End'],"%d %b %Y")
                #print('----------- Error of ', Error.days,' days in ',v1['Duration'],' day Dasha')
            count = count+1


def tracer():
    print('tracer')
    print(p21.ChartType)
    
#print('p21utilsR imported')
