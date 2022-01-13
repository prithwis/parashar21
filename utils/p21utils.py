# 𝓟𝓪𝓻𝓪𝓼𝓱𝓪𝓻21
# 𝓟𝓻𝓲𝓽𝓱𝔀𝓲𝓼 𝓜𝓾𝓴𝓮𝓻𝓳𝓮𝓮
# --------------------------------------------------
# Global Variables
import p21
# --------------------------------------------------
#Utility functions 

import pandas as pd
import dateutil
import matplotlib.pyplot as plt
import math
import math
import numbers
import string_utils

import json

from docx import Document
from docx.shared import Inches
from datetime import datetime
from datetime import timedelta

import swisseph as swe
import pytz

# --------------------------------------------------
#Configure Swiss Ephemeris

def C01_configSWE():
    
    swe.version
    #Constants and Flags necessary for this application are set here
    #
    swe.set_ephe_path('/content/ephe') # set path to ephemeris files
    # Calendar : Julian or Gregorian
    SE_GREG_CAL = 1
    p21.gregflag = SE_GREG_CAL
    #
    # Ayanamsha type : Lahiri = 1
    swe.set_sid_mode(1)  # Lahiri Aynamsha
    #
    # whether speed will be calculated along with position of planet
    SEFLG_SPEED = int(256)
    p21.iflag = SEFLG_SPEED
    #hsys = P, Placidus
    #ascii P = 080
    p21.hsysP =  bytes('P', 'utf-8')
    
# --------------------------------------------------
#
def C02_parsePersonData(P):
    #global chart
    global LBDY, LBDM,LBDD,LBTh,LBTm,LBTs,tZone,natalLAT,natalLON
    
    LBDY = P['DoB_Year']     # Year of Birth, Local Time
    LBDM = P['DoB_Mon']      # Month of Birth, Local Time
    LBDD = P['DoB_Day']      # Day of Birth, Local Time

    LBTh = int(P['DoB_Time'].split(":")[0])                           # Hour of Birth, Local Time, 24 Hour Clock
    LBTm = int(P['DoB_Time'].split(":")[1])                          # Minute of Birth, Local Time
    LBTs = 0                           # Second of Birth, Local Time

    tZone = P['TZ_OffHours']                         # Time Zone Offset. East of Greenwich is +ve
    #print('ppd', LBDY, LBDM,LBDD,LBTh,LBTm,LBTs,tZone)
    #print(P['Gender'])
    natalLAT = P['PoB_Lat']
    natalLON = P['PoB_Lon']
    ck = P['Gender']+str(P['DoB_Year'])+str(P['DoB_Mon'])+str(P['DoB_Day'])+P['DoB_Time']+str(P['TZ_OffHours'])+str(P['PoB_Lat'])+str(P['PoB_Lon'])
    p21.chart = {
        "pid":{
            "tags" : [P['Voc1_Cat'],P['Voc2_Cat'],P['Voc3_Cat']],
            "ck" : ck,
            "name": string_utils.shuffle(P['Name'])
        }
    }
    
# --------------------------------------------------

def C03_convertDates():
    global natalUT, ayanamsha
    # Date & Time of Birth
    #Time Zone Conversion
    UTCdt = swe.utc_time_zone(LBDY,LBDM,LBDD,LBTh,LBTm,LBTs,tZone)
    #print(UTCdt)
    # Date & Time of Birth
    # Convert UTC date to Julian Date
    JD = swe.utc_to_jd(UTCdt[0],UTCdt[1],UTCdt[2],UTCdt[3],UTCdt[4],0,p21.gregflag) # Second Value = 0
    #print (JD[0], JD[1])
    # JD[0] is Ephemeris Time
    # JD[1] is Universal Time 
    natalUT = JD[1]
    #print(natalUT, 'Julian Date of DTofB in Universal Time')
    #Ayanamsha
    ayanamsha = swe.get_ayanamsa(natalUT)
    #print('Lahiri Ayanamsha :', ayanamsha)
    
# --------------------------------------------------
#
def C04_calculateGrahaPositions():
    global bLon,bRet,P
    #Position of Planets
    #
    # body 0 = Sun, 1 = Moon, 2 = Merc, 3 = Ven, 4 = Mars, 5 = Jup, 6 = Sat, 11 = True Node (Rahu)
    bLon = []
    bRet = []
    for body in [0,1,2,3,4,5,6,11]:
        pData = swe.calc_ut(natalUT,body,p21.iflag)
        bLon.append(pData[0][0])
        if pData[0][3] >= 0:
            bRet.append(False)
        else :
            bRet.append(True)
    #for ix in range(len(bLon)) : print(ix, bLon[ix], bRet[ix])
    #-------------------------------------------------------------------
    #House Position and Ascendants
    P = swe.houses(natalUT,natalLAT,natalLON,p21.hsysP)
    
    #P = swe.houses_ex(natalUT,natalLAT,natalLON,hsysP)
    #for ix in range(len(P[0])) : print(ix, P[0][ix])
    #print('--')
    #for ix in range(len(P[1])) : print(ix, P[1][ix])

# --------------------------------------------------

def C05_buildGLonGRet():
    #global GLon, GRet
    p21.GLon = {
    'La':p21Sub(P[1][0],ayanamsha),
    'Su':p21Sub(bLon[0],ayanamsha),
    'Mo':p21Sub(bLon[1],ayanamsha),
    'Ma':p21Sub(bLon[4],ayanamsha),
    'Me':p21Sub(bLon[2],ayanamsha),
    'Ju':p21Sub(bLon[5],ayanamsha),
    'Ve':p21Sub(bLon[3],ayanamsha),
    'Sa':p21Sub(bLon[6],ayanamsha),
    'Ra':p21Sub(bLon[7],ayanamsha),
    'Ke':p21Sub(p21Sub(bLon[7],ayanamsha),180)
    }
    
    p21.GRet = {
    'La':False,
    'Su':False,
    'Mo':False,
    'Ma':bRet[4],
    'Me':bRet[2],
    'Ju':bRet[5],
    'Ve':bRet[3],
    'Sa':bRet[6],
    'Ra':False,
    'Ke':False
    }
    p21.GLonRet = {
        'GLon': p21.GLon,
        'GRet': p21.GRet
    }

    

# --------------------------------------------------
# Subtracts angles (eg Ayanamsha) from Longitudes 
# Required for conversion from Tropical (Western) to Sidereal (Indian) zodiac
#
def p21Sub (x,y):
    retVal = round(x - y,3)
    if retVal < 0:
        return retVal + 360
    else:
        return retVal
# --------------------------------------------------
# append one ditionary after another

def appendDict(d1,d2):
    for i in range(len(d2)):
        k = list(d2.keys())[i]
        v = list(d2.values())[i]
        d1[k] = v


# --------------------------------------------------
#converting dict{ionary} to list and back

d2l = lambda dic: [(k, v) for (k, v) in dic.items()]
l2d = lambda lis: dict(lis)

# --------------------------------------------------

#converts a Rashi number to the Rashi name
#
RashiName = ["RashiName","Ari","Tau","Gem","Can","Leo","Vir","Lib","Sco","Sag","Cap","Acq","Pis"]
def RashiN2A(n):
    return RashiName[n]

# --------------------------------------------------

#Determines and stores the Bhava in list BhavN
#BhavN[1] has Rashi Number corresponding to First Bhav
#BhavA[2] has the Rashi name corresponding to Second Bhava
#
def C10_DetermineBhavs():

    p21.BhavN = [p21.BoL]
    for ix in range(1,13):
        if ix == 1:
            p21.BhavN.append(p21.GRashiN['La'])
        else:
            N = p21.BhavN[ix-1]+1
            if N > 12:
                N = N - 12
            p21.BhavN.append(N)
        
    
    p21.BhavA = list(map(lambda x : RashiN2A(x) if isinstance(x, numbers.Integral) else p21.BoL,p21.BhavN))
    
    p21.BhavNBhavA = {
        'BhavN': p21.BhavN,
        'BhavA': p21.BhavA
    }

# --------------------------------------------------

#Determines the Lord of each Bhava
#Determines the Rashi number where each Lord resides
#Determines the Rashi name where each Lord resides
#
def C11_DetermineLord():

    p21.LordOf = {"Ari":"Ma","Tau":"Ve","Gem":"Me","Can":"Mo","Leo":"Su","Vir":"Me","Lib":"Ve","Sco":"Ma","Sag":"Ju","Cap":"Sa","Acq":"Sa","Pis":"Ju"}
    #print(LordOf)
    p21.Lord = list(map(lambda x : p21.LordOf[RashiN2A(x)] if isinstance(x, numbers.Integral) else p21.BoL, p21.BhavN))
    #print("Lord : ", p21.Lord)
    p21.LordRashiN = list(map(lambda x : p21.GRashiN[x] if x != p21.BoL else p21.BoL, p21.Lord))
    #print("Rashi # of Lord : ",p21.LordRashiN)
    p21.LordRashiA = list(map(lambda x : p21.GRashiA[x] if x != p21.BoL else p21.BoL, p21.Lord))
    #print("Rashi Name of Lord :",p21.LordRashiA)

# --------------------------------------------------

#Determines the Bhava where each Graha is Lord
#Grahas other than Sun and Moon are Lords of Two Bhava
#Since index returns only the first position where an instance is found
#A simple index() does not return the second position
# https://stackoverflow.com/questions/22267241/how-to-find-the-index-of-the-nth-time-an-item-appears-in-a-list


    p21.GrahaLordBhav = {}
    for G in ('Su','Mo','Ma','Me','Ju','Ve','Sa'):
            L = [i for i, n in enumerate(p21.Lord) if n == G]
            p21.GrahaLordBhav[G] = L
    #print('Bhav where Graha is Lord : ',p21.GrahaLordBhav)
    
    p21.LordInfo = {
     'Lord' : p21.Lord,
     'LordRashiN' : p21.LordRashiN,
     'LordRashiA' : p21.LordRashiA,
     'GrahaLordBhav' : p21.GrahaLordBhav
    }


def C12_BhavOfGraha_Lord():

#Determines the Bhava where Each planet resides
#Lagna always resides in First Bhava
#
    p21.GrahaBhava ={"La":1}
    for G in ('Su','Mo','Ma','Me','Ju','Ve','Sa','Ra','Ke'):
        p21.GrahaBhava[G] = p21.BhavN.index(p21.GRashiN[G])
    print (p21.GrahaBhava)

#Determines the Bhav where each Lord resides
#

    p21.LordBhav =[p21.BoL]
    for L in range (1,13):
        p21.LordBhav.append(p21.BhavN.index(p21.GRashiN[p21.Lord[L]]))
    print(p21.LordBhav)

    p21.BhavOfGraha_LordInfo = {
     'GrahaBhava' : p21.GrahaBhava,
     'LordBhav' : p21.LordBhav
    }

# --------------------------------------------------
#
# Generates the text that appears in a chart image
#
def generateChartTxt():
    global txt
    txt = ['']*13
    for g,r in p21.GRashiN.items():
        if r == 1:
            txt[1]= txt[1]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 2:
            txt[2]= txt[2]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 3:
            txt[3]= txt[3]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 4:
            txt[4]= txt[4]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 5:
            txt[5]= txt[5]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 6:
            txt[6]= txt[6]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 7:
            txt[7]= txt[7]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 8:
            txt[8]= txt[8]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 9:
            txt[9]= txt[9]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 10:
            txt[10]= txt[10]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 11:
            txt[11]= txt[11]+' '+g+('/R' if p21.GRet[g] else '')
        if r == 12:
            txt[12]= txt[12]+' '+g+('/R' if p21.GRet[g] else '')
            
    for i in range(1,13):
        if len(txt[i]) == 0:
            txt[i] = '*'
#print(txt[1],txt[2],txt[3],txt[4],txt[5],txt[6],txt[7],txt[8],txt[9],txt[10],txt[11],txt[12],)

# --------------------------------------------------
# Plots the chart in the Bengal Format
#
def R12B_drawChart_Bengal():

    #ChartFile = p21.pName+p21.ChartType+'.png'
    #ChartFile = 'CurrentChart.png'
    generateChartTxt()

    #id = ChartType+'\n'+pName+'\n'+pDate+'\n'+pTime+'\n'+pPlace
    id = p21.ChartType+'\n'+p21.pName
    if p21.ChartType == 'Rashi':
        ChartColour = 'orange'
        ChartFile = 'RashiChart.png'
    else:
        ChartColour = 'olive'
        ChartFile = 'NavamsaChart.png'
    
        
    #plt.figure(figsize=(7,7))
    plt.figure(figsize=(7,7),facecolor=ChartColour)
    
    plt.axis('off')

    # draw vertical line 
    plt.plot([30, 30], [0, 90], 'k-', lw=2)
    plt.plot([60, 60], [0, 90], 'k-', lw=2)

    # draw horizontal line 
    plt.plot([0, 90], [30, 30], 'k-', lw=2)
    plt.plot([0, 90], [60, 60], 'k-', lw=2)

    #draw diagonal lines
    plt.plot([60,90],[60,90], 'k-', lw=2)
    plt.plot([0,30],[90,60], 'k-', lw=2)
    plt.plot([0,30],[90,60], 'k-', lw=2)
    plt.plot([0,30],[0,30], 'k-', lw=2)
    plt.plot([60,90],[30,0], 'k-', lw=2)

    plt.text(32, 38, id, fontsize=12)

    plt.text(32, 82, txt[1], fontsize=12)
    plt.text(8, 82, txt[2], fontsize=12)
    plt.text(2, 62, txt[3], fontsize=12)
    plt.text(2, 45, txt[4], fontsize=12)
    plt.text(2, 25, txt[5], fontsize=12)
    plt.text(8, 5, txt[6], fontsize=12)
    plt.text(32, 5, txt[7], fontsize=12)
    plt.text(62, 5, txt[8], fontsize=12)
    plt.text(68, 25, txt[9], fontsize=12)
    plt.text(68, 45, txt[10], fontsize=12)
    plt.text(68, 62, txt[11], fontsize=12)
    plt.text(62, 82, txt[12], fontsize=12)

    
    #plt.savefig("CurrentChart.png", bbox_inches='tight')
    plt.savefig(ChartFile, bbox_inches='tight')
    #plt.show()


# --------------------------------------------------

# Convert the longitude of a Graha into its Rashi of residence
#
# Navamsa is different
# http://www.oocities.org/talk2astrologer/LearnAstrology/Details/Navamsa.html
#
# the logic for placing a Graha in a 'Navamsa' Rashi is as follows
# starting from Longitude 0 [ 'Normal' Aries / Mesha start point]
# upto Longitude 360, we divide the longitudes in 108 partitions with each 
# partition being 3.33333 degrees. These 108 partitions are now numbered 
# sequentially, except that once we reach 12, the next partition is 1, not
# 13. So we have 9 sequences running from 1 - 12 each. The number corresponding
# to the partition where the Graha falls is the Navamsa Rashi of the Graha


def Long2Rashi(x):
    
    if p21.ChartType == 'Rashi':
        RashiNumber = math.floor(x[1]/30)+1
        return x[0],RashiNumber
    if p21.ChartType == 'Navamsa':
        N1 = math.floor(x[1]/3.333333)+1
        N2 = N1%12
        if N2 == 0:
            RashiNumber = 12
        else:
            RashiNumber = N2
        return x[0],RashiNumber

# --------------------------------------------------
#Defines the Horoscope in terms of locating planets in their Rashis

def R11_LocateGrahaInRashi():
    
    
    #print('Loc',p21.ChartType)
    p21.GRashiN = l2d(list(map(lambda x : Long2Rashi(x), d2l(p21.GLon))))
    #print(p21.GRashiN)

    p21.GRashiA = {}
    for k,v in p21.GRashiN.items():
        p21.GRashiA[k] = RashiN2A(v)
        
    #print(p21.GRashiA)
    
# --------------------------------------------------

def R01_CreateReportDoc(cqs):
      

    p21.document = Document()
    section = p21.document.sections[0]
    header = section.header
    footer = section.footer
    
    header01 = header.paragraphs[0]
    header01.text = "Parashar21 | Khona21 MongoDB database"

    now = datetime.now(pytz.timezone('Asia/Kolkata'))
    #print(now.strftime("%d %b %Y"))

    footer01 = footer.paragraphs[0]
    footer01.text = "Printed on : "+now.strftime("%d %b %Y")+" | http://parashar21.blogspot.com | https://github.com/prithwis/parashar21"

    #heading_1 = p21.pName +" >>> "+p21.ChartType
    #heading_1 = ChartType+" Chart of "+pName
    heading_1 = "Selected Charts"  #p21.ReportFile
    p21.document.add_heading(heading_1, 0)
    #p21.document.add_picture('./Saraswati.png', width=Inches(4.25))
    heading_2 = cqs                         # current query string
    p21.document.add_heading(heading_2, level=3)

    
    #p21.document.add_page_break()
    #return document
    
# --------------------------------------------------

def R13B_ListPositions(text,L):
    # 
    L1 = L[1:]
    D = text+' : '
    for c,e in enumerate(L1,1):
        D = D+' '+str(c)+':'+e+' '
    return(D)
        

def R13A_ShowTrueDict(desc,tDict):
    # creates a string with Keys that have Value True
    
    D = desc+' : '
    for (k,v) in tDict.items():
        if v :
            D = D+' '+k
    return(D)


def tracer():
    print('tracer')
    print(p21.ChartType)
    
print('p21utils imported')
# ------------------------------------------------